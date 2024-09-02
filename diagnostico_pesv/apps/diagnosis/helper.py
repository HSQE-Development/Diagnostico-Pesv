from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.table import _Cell
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
import openpyxl
from openpyxl.chart import BarChart, Reference
import pandas as pd
import os
import base64
from tempfile import NamedTemporaryFile
from docx2pdf import convert


def apply_bullets(paragraph):
    """Aplica viñetas al párrafo usando XML."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")  # Utiliza el ID de numeración predeterminado
    numPr.append(numId)
    pPr.append(numPr)


def add_title(doc: Document, title_text: str):
    # Agregar un párrafo con el estilo de título
    title_paragraph = doc.add_paragraph(title_text, style="Body Text")


def align_cell_text(cell, horizontal="center", vertical="center"):
    """
    Aligns text in a cell both horizontally and vertically.

    :param cell: The cell to align.
    :param horizontal: Horizontal alignment ('left', 'center', 'right').
    :param vertical: Vertical alignment ('top', 'center', 'bottom').
    """
    # Horizontal alignment
    if horizontal == "left":
        cell.paragraphs[0].alignment = 0  # Left-align
    elif horizontal == "center":
        cell.paragraphs[0].alignment = 1  # Center-align
    elif horizontal == "right":
        cell.paragraphs[0].alignment = 2  # Right-align

    # Vertical alignment
    cell_properties = cell._element.get_or_add_tcPr()

    # Ensure the namespace is defined
    cell_properties.set(qn("w:tcPr"), "")

    v_align = cell_properties.find(qn("w:vAlign"))

    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        cell_properties.append(v_align)

    if vertical == "top":
        v_align.set(qn("w:val"), "top")
    elif vertical == "center":
        v_align.set(qn("w:val"), "center")
    elif vertical == "bottom":
        v_align.set(qn("w:val"), "bottom")


def trim_merged_cell(cell):
    """
    Trims the text of a merged cell and clears the text in other merged cells.

    :param cell: The cell to trim.
    """
    # Trim text in the current cell
    cell.text = cell.text.strip()

    # Find all cells in the merged range and clear their text
    for row in cell._element.getparent().iterchildren():
        for c in row.iterchildren():
            if c is not cell._element and c.tag == qn("w:tc"):
                c.text = ""


def hex_to_rgb(hex_color):
    """
    Convert hexadecimal color to RGB tuple.

    :param hex_color: Hexadecimal color string, e.g., 'FF0000'.
    :return: Tuple (R, G, B).
    """
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def rgb_to_hex(rgb_color):
    """
    Convert RGB tuple to hexadecimal color string.

    :param rgb_color: Tuple (R, G, B).
    :return: Hexadecimal color string, e.g., 'FF0000'.
    """
    return "".join(f"{c:02X}" for c in rgb_color).upper()


def luminance(rgb_color):
    """
    Calculate the luminance of an RGB color to determine if it's dark or light.

    :param rgb_color: Tuple (R, G, B).
    :return: Luminance value.
    """
    r, g, b = [x / 255.0 for x in rgb_color]
    # Apply the luminance formula
    return 0.2126 * (r**2.2) + 0.7152 * (g**2.2) + 0.0722 * (b**2.2)


def set_cell_background_color(cell, color):
    """
    Set the background color of a cell using either RGB or hexadecimal color values.
    Automatically adjusts the text color based on the background color's luminance.

    :param cell: The cell to modify.
    :param color: RGB color value as a tuple, e.g., (255, 0, 0) for red,
                  or a hexadecimal string, e.g., 'FF0000' for red.
    """
    # Check if the color is in RGB tuple format
    if isinstance(color, tuple):
        hex_color = rgb_to_hex(color)
        rgb_color = color
    elif isinstance(color, str):
        hex_color = color.lstrip("#").upper()
        rgb_color = hex_to_rgb(hex_color)
    else:
        raise ValueError("Color must be either an RGB tuple or a hexadecimal string")

    # Ensure hex_color is 6 characters long
    hex_color = hex_color.zfill(6)

    # Define the XML namespace for background color
    cell_properties = cell._element.get_or_add_tcPr()
    shd = cell_properties.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        cell_properties.append(shd)

    # Set the fill color
    shd.set(qn("w:fill"), hex_color)

    # Determine text color based on background color's luminance
    luminance_value = luminance(rgb_color)
    text_color = RGBColor(255, 255, 255) if luminance_value < 0.5 else RGBColor(0, 0, 0)

    # Apply text color
    p = cell.paragraphs[0]
    run = p.add_run()
    run.font.color.rgb = text_color


def set_cell_text_color(cell, color=None):
    """
    Set the text color of a cell. If no color is specified, the default is white.

    :param cell: The cell to modify.
    :param color: RGB color value as a tuple, e.g., (255, 0, 0) for red,
                  or a hexadecimal string, e.g., 'FF0000' for red.
                  If None, sets the text color to white.
    """
    # Default to white color if no color is specified
    if color is None:
        color = (255, 255, 255)  # White

    # Check if the color is in hexadecimal format or RGB tuple
    if isinstance(color, str):
        rgb_color = hex_to_rgb(color)
    elif isinstance(color, tuple):
        rgb_color = color
    else:
        raise ValueError("Color must be either an RGB tuple or a hexadecimal string")

    # Create an RGBColor object
    text_color = RGBColor(*rgb_color)

    # Apply the color to all runs in the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = text_color


def replace_text_in_paragraph(paragraph, search_text, replace_text):
    if search_text in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if search_text in item.text:
                item.text = item.text.replace(search_text, replace_text)
                if search_text == "{{CONSULTOR_NOMBRE}}":
                    item.bold = True
                if search_text == "{{MISIONALIDAD_NAME}}":
                    item.bold = True
                if search_text == "{{MISIONALIDAD_ID}}":
                    item.bold = True
                if search_text == "{{NIVEL_PESV}}":
                    item.bold = True
                if search_text == "{{QUANTITY_VEHICLES}}":
                    item.bold = True
                if search_text == "{{QUANTITY_DRIVERS}}":
                    item.bold = True


def replace_text_in_table(table, search_text, replace_text):
    for row in table.rows:
        for cell in row.cells:
            replace_text_in_paragraphs(cell.paragraphs, search_text, replace_text)


def replace_text_in_paragraphs(paragraphs, search_text, replace_text):
    for paragraph in paragraphs:
        replace_text_in_paragraph(paragraph, search_text, replace_text)


def replace_placeholders_in_document(doc: Document, placeholders: dict):
    # Reemplaza los marcadores en todos los párrafos
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            replace_text_in_paragraph(paragraph, placeholder, replacement)


def format_nit(nit):
    # Convierte el valor a cadena y elimina caracteres no numéricos
    nit_string = "".join(filter(str.isdigit, str(nit)))

    # Verifica que la longitud sea de 10 dígitos
    if len(nit_string) != 10:
        raise ValueError("El NIT debe tener 10 dígitos.")

    # Aplica el formato XXXXXXXXX-X
    formatted_nit = f"{nit_string[:9]}-{nit_string[9]}"

    return formatted_nit


def get_current_month_and_year():
    # Obtiene la fecha y hora actuales
    now = datetime.now()

    # Traducciones de nombres de meses al español
    meses_en_espanol = [
        "Enero",
        "Febrero",
        "Marzo",
        "Abril",
        "Mayo",
        "Junio",
        "Julio",
        "Agosto",
        "Septiembre",
        "Octubre",
        "Noviembre",
        "Diciembre",
    ]

    # Obtiene el mes en número (1-12) y lo traduce al nombre del mes en español
    numero_mes = now.month
    nombre_mes = meses_en_espanol[numero_mes - 1]
    año_actual = now.year

    return nombre_mes, año_actual


def insert_table_after_placeholder(doc: Document, placeholder: str, table_data: list):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = "Table Grid"
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_data
            table._element.getparent().insert(index + 1, table._element)
            return


def insert_table_after_placeholder(
    doc: Document,
    placeholder: str,
    fecha: str,
    empresa: str,
    nit: str,
    actividades: str,
    vehicle_questions: list,
    fleet_data: list,
    driver_questions,
    driver_data,
    com_size,
    segment,
    contact,
    certification,
    ciius,
):
    """
    Inserta una tabla en el documento justo después del párrafo que contiene el placeholder.

    :param doc: Documento Word en el que se insertará la tabla.
    :param placeholder: Palabra clave para encontrar la ubicación de la tabla.
    :param fecha: Fecha de elaboración.
    :param empresa: Nombre de la empresa.
    :param nit: NIT de la empresa.
    :param actividades: Actividades de la empresa.
    :param vehicle_questions: Lista de preguntas sobre vehículos.
    :param fleet_data: Lista de datos de la flota.
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Insertar la tabla después del párrafo que contiene el placeholder
            index = paragraph._element.getparent().index(paragraph._element)

            # Crear la tabla con el formato especificado
            table = doc.add_table(rows=1, cols=12)
            table.style = "Table Grid"

            # Configurar el ancho de las columnas (opcional)
            for col in table.columns:
                col.width = Inches(2)

            # Agregar encabezado para "CARACTERIZACION DE LA EMPRESA"
            heading_row = table.rows[0].cells
            heading_row[0].text = "CARACTERIZACION DE LA EMPRESA"
            heading_row[0].merge(heading_row[11])

            # Agregar fila con fecha y empresa
            row_cells = table.add_row().cells
            row_cells[0].text = "Fecha de elaboración"
            row_cells[0].merge(row_cells[1])
            row_cells[2].text = fecha
            row_cells[2].merge(row_cells[5])
            row_cells[6].text = "Empresa"
            row_cells[6].merge(row_cells[7])
            row_cells[8].text = empresa
            row_cells[8].merge(row_cells[11])
            for cell in row_cells:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)
            # Agregar fila con NIT y actividades
            row_cells = table.add_row().cells
            row_cells[0].text = "Nit"
            row_cells[0].merge(row_cells[1])
            row_cells[2].text = nit
            row_cells[2].merge(row_cells[5])
            row_cells[6].text = "Actividades"
            row_cells[6].merge(row_cells[7])
            activities_cell = row_cells[8]
            activities_cell.add_paragraph()
            # Añadir cada Ciiu como ítem en la lista
            for ciiu in ciius.all():
                bullet_paragraph = activities_cell.add_paragraph()
                bullet_paragraph.text = f"{ciiu.code} - {ciiu.name}"
                apply_bullets(bullet_paragraph)
            row_cells[8].merge(row_cells[11])
            for cell in row_cells:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)
            # Agregar fila con NIT y actividades
            row_cells = table.add_row().cells
            row_cells[0].text = "Tamaño de la empresa"
            row_cells[0].merge(row_cells[1])
            row_cells[2].text = com_size
            row_cells[2].merge(row_cells[5])
            row_cells[6].text = "Segmento al que pertenece"
            row_cells[6].merge(row_cells[7])
            row_cells[8].text = segment
            row_cells[8].merge(row_cells[11])
            for cell in row_cells:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)
            # Agregar fila con NIT y actividades
            row_cells = table.add_row().cells
            row_cells[0].text = "Contacto"
            row_cells[0].merge(row_cells[1])
            row_cells[2].text = contact
            row_cells[2].merge(row_cells[5])
            row_cells[6].text = "Certificaciones adquiridas (Normas ISO)"
            row_cells[6].merge(row_cells[7])
            row_cells[8].text = certification or "NINGUNA"
            row_cells[8].merge(row_cells[11])

            for cell in row_cells:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)

            # Agregar fila para Flota de vehículos
            flota_header_row = table.add_row().cells
            flota_header_row[0].text = "Flota de vehículos"
            flota_header_row[0].merge(flota_header_row[11])
            align_cell_text(flota_header_row[0], "left")
            set_cell_background_color(flota_header_row[0], "2f4858")
            set_cell_text_color(flota_header_row[0])

            flota_rows_row = table.add_row().cells
            flota_rows_row[0].text = "FLOTA DE VEHICULOS AUTOMOTORES"
            flota_rows_row[0].merge(flota_rows_row[4])
            flota_rows_row[5].text = "Cantidad Propios"
            flota_rows_row[6].text = "Cantidad Terceros"
            flota_rows_row[7].text = "Cantidad Arrendados"
            flota_rows_row[8].text = "Cantidad Contratistas"
            flota_rows_row[9].text = "Cantidad Intermediación"
            flota_rows_row[10].text = "Cantidad Leasing"
            flota_rows_row[11].text = "Cantidad Renting"
            for cell in flota_rows_row:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)

            # Variables para almacenar los totales
            total_propio = 0
            total_tercero = 0
            total_arrendado = 0
            total_contratista = 0
            total_intermediacion = 0
            total_leasing = 0
            total_renting = 0

            # Insertar datos de flota
            for vehicle_question in vehicle_questions:
                row_cells = table.add_row().cells
                row_cells[0].text = vehicle_question.name
                row_cells[0].merge(row_cells[4])
                align_cell_text(row_cells[0], "left")
                set_cell_background_color(row_cells[0], "2f4858")
                set_cell_text_color(row_cells[0])
                fleet = next(
                    (
                        f
                        for f in fleet_data
                        if f.vehicle_question.id == vehicle_question.id
                    ),
                    None,
                )
                quantity_propio = fleet.quantity_owned if fleet else 0
                quantity_tercero = fleet.quantity_third_party if fleet else 0
                quantity_arrendado = fleet.quantity_arrended if fleet else 0
                quantity_contratista = fleet.quantity_contractors if fleet else 0
                quantity_intermediacion = fleet.quantity_intermediation if fleet else 0
                quantity_leasing = fleet.quantity_leasing if fleet else 0
                quantity_renting = fleet.quantity_renting if fleet else 0

                row_cells[5].text = str(quantity_propio)
                row_cells[6].text = str(quantity_tercero)
                row_cells[7].text = str(quantity_arrendado)
                row_cells[8].text = str(quantity_contratista)
                row_cells[9].text = str(quantity_intermediacion)
                row_cells[10].text = str(quantity_leasing)
                row_cells[11].text = str(quantity_renting)
                for cell in row_cells:
                    align_cell_text(cell)

                # Sumar cantidades a los totales
                total_propio += quantity_propio
                total_tercero += quantity_tercero
                total_arrendado += quantity_arrendado
                total_contratista += quantity_contratista
                total_intermediacion += quantity_intermediacion
                total_leasing += quantity_leasing
                total_renting += quantity_renting

                total = (
                    total_propio
                    + total_tercero
                    + total_arrendado
                    + total_contratista
                    + total_intermediacion
                    + total_leasing
                    + total_renting
                )

            # Agregar fila con los totales
            total_row = table.add_row().cells
            total_row[0].text = "Total Vehiculos".upper()
            total_row[0].merge(total_row[4])
            align_cell_text(total_row[0], "left")
            set_cell_background_color(total_row[0], "2f4858")
            set_cell_text_color(total_row[0])
            total_row[5].text = str(total)
            total_row[5].merge(total_row[6])
            total_row[5].merge(total_row[7])
            total_row[5].merge(total_row[8])
            total_row[5].merge(total_row[9])
            total_row[5].merge(total_row[10])
            total_row[5].merge(total_row[11])

            # Agregar fila para Conductores
            driver_header_row = table.add_row().cells
            driver_header_row[0].text = (
                "Personas que conducen con fines misionales".upper()
            )
            driver_header_row[0].merge(driver_header_row[8])
            driver_header_row[9].text = "Cantidad"
            driver_header_row[9].merge(driver_header_row[11])
            for cell in driver_header_row:
                align_cell_text(cell, "left")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)
            total_conductores = 0
            # Datos de conductores
            for driver_question in driver_questions:
                driver_data_row = table.add_row().cells
                driver_data_row[0].text = driver_question.name
                driver_data_row[0].merge(driver_data_row[8])
                align_cell_text(driver_data_row[0], "left")
                set_cell_background_color(driver_data_row[0], "2f4858")
                set_cell_text_color(driver_data_row[0])
                driver = next(
                    (
                        f
                        for f in driver_data
                        if f.driver_question.id == driver_question.id
                    ),
                    None,
                )
                quantity = driver.quantity if driver else 0
                driver_data_row[9].text = str(quantity)
                driver_data_row[9].merge(driver_data_row[11])
                for cell in driver_data_row:
                    align_cell_text(cell)
                total_conductores += quantity
            # Agregar fila con los totales
            total_driver_row = table.add_row().cells
            total_driver_row[0].text = "Total Conductores".upper()
            total_driver_row[0].merge(total_driver_row[8])
            align_cell_text(total_driver_row[0], "left")
            set_cell_background_color(total_driver_row[0], "2f4858")
            set_cell_text_color(total_driver_row[0])
            total_driver_row[9].text = str(total_conductores)
            total_driver_row[9].merge(total_driver_row[11])
            # Mover la tabla a la posición deseada
            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_tables_for_companies(
    doc: Document,
    placeholder: str,
    companies: list,
    fecha,
    vehicle_questions: list,
    driver_questions: list,
    Fleet,
    Driver,
    diagnosis,
):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            for company_data in companies:
                company = company_data["company"]
                count_size = company_data["count_size"]
                total_owned = company_data["total_owned"]
                total_leasing = company_data["total_leasing"]
                total_renting = company_data["total_renting"]
                total_quantity_driver = company_data["total_quantity_driver"]

                title_paragraph = doc.add_paragraph()
                title_paragraph.add_run(f"{company.name.upper()}").bold = True
                title_paragraph.alignment = 1  # Centrar el título
                # Insertar el título en el lugar correcto
                paragraph._element.getparent().insert(
                    index + 1, title_paragraph._element
                )
                index += 1  # Ajustar el índice para el siguiente elemento

                # Insertar un párrafo vacío para separar el título de la tabla
                doc.add_paragraph()  # Opcional, para mayor claridad visual
                paragraph._element.getparent().insert(
                    index + 1, doc.paragraphs[-1]._element
                )
                index += 1  # Ajustar el índice para el siguiente elemento

                # Crear la tabla con el formato especificado
                table = doc.add_table(rows=1, cols=12)
                table.style = "Table Grid"

                # Encabezado para "CARACTERIZACION DE LA EMPRESA"
                heading_row = table.rows[0].cells
                heading_row[0].text = (
                    f"CARACTERIZACION DE LA EMPRESA - {company.name.upper()}"
                )
                heading_row[0].merge(heading_row[11])
                for cell in heading_row:
                    align_cell_text(cell, "center", "center")

                row_cells = table.add_row().cells
                row_cells[0].text = "Fecha de elaboración"
                row_cells[0].merge(row_cells[1])
                row_cells[2].text = fecha
                row_cells[2].merge(row_cells[5])
                row_cells[6].text = "Razón Social"
                row_cells[6].merge(row_cells[7])
                row_cells[8].text = company.name.upper()
                row_cells[8].merge(row_cells[11])
                for cell in row_cells:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)
                    # Agregar fila con NIT y actividades
                row_cells = table.add_row().cells
                row_cells[0].text = "Nit"
                row_cells[0].merge(row_cells[1])
                row_cells[2].text = company.nit
                row_cells[2].merge(row_cells[5])
                row_cells[6].text = "Actividades"
                row_cells[6].merge(row_cells[7])
                activities_cell = row_cells[8]
                activities_cell.add_paragraph()
                # Añadir cada Ciiu como ítem en la lista
                for ciiu in company.ciius.all():
                    bullet_paragraph = activities_cell.add_paragraph()
                    bullet_paragraph.text = f"{ciiu.code} - {ciiu.name}"
                    apply_bullets(bullet_paragraph)
                row_cells[8].merge(row_cells[11])
                for cell in row_cells:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)
                # Agregar fila con NIT y actividades
                row_cells = table.add_row().cells
                row_cells[0].text = "Tamaño de la empresa"
                row_cells[0].merge(row_cells[1])
                row_cells[2].text = count_size.name.upper()
                row_cells[2].merge(row_cells[5])
                row_cells[6].text = "Segmento al que pertenece"
                row_cells[6].merge(row_cells[7])
                row_cells[8].text = company.segment.name.upper()
                row_cells[8].merge(row_cells[11])
                for cell in row_cells:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)
                # Agregar fila con NIT y actividades
                row_cells = table.add_row().cells
                row_cells[0].text = "Contacto"
                row_cells[0].merge(row_cells[1])
                row_cells[2].text = company.dependant
                row_cells[2].merge(row_cells[5])
                row_cells[6].text = "Certificaciones adquiridas (Normas ISO)"
                row_cells[6].merge(row_cells[7])
                row_cells[8].text = company.acquired_certification or "NINGUNA"
                row_cells[8].merge(row_cells[11])

                for cell in row_cells:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)

                # Agregar fila para Flota de vehículos
                flota_header_row = table.add_row().cells
                flota_header_row[0].text = "Flota de vehículos"
                flota_header_row[0].merge(flota_header_row[11])
                align_cell_text(flota_header_row[0], "left")
                set_cell_background_color(flota_header_row[0], "2f4858")
                set_cell_text_color(flota_header_row[0])

                flota_rows_row = table.add_row().cells
                flota_rows_row[0].text = "FLOTA DE VEHICULOS AUTOMOTORES"
                flota_rows_row[0].merge(flota_rows_row[4])
                flota_rows_row[5].text = "Cantidad Propios"
                flota_rows_row[6].text = "Cantidad Terceros"
                flota_rows_row[7].text = "Cantidad Arrendados"
                flota_rows_row[8].text = "Cantidad Contratistas"
                flota_rows_row[9].text = "Cantidad Intermediación"
                flota_rows_row[10].text = "Cantidad Leasing"
                flota_rows_row[11].text = "Cantidad Renting"
                for cell in flota_rows_row:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)

                fleet_data = Fleet.objects.filter(
                    diagnosis_counter__company=company,
                    diagnosis_counter__diagnosis=diagnosis,
                )
                driver_data = Driver.objects.filter(
                    diagnosis_counter__company=company,
                    diagnosis_counter__diagnosis=diagnosis,
                )
                # Variables para almacenar los totales
                total_propio = 0
                total_tercero = 0
                total_arrendado = 0
                total_contratista = 0
                total_intermediacion = 0
                total_leasing = 0
                total_renting = 0
                # Insertar datos de flota
                for vehicle_question in vehicle_questions:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vehicle_question.name
                    row_cells[0].merge(row_cells[4])
                    align_cell_text(row_cells[0], "left")
                    set_cell_background_color(row_cells[0], "2f4858")
                    set_cell_text_color(row_cells[0])
                    fleet = next(
                        (
                            f
                            for f in fleet_data
                            if f.vehicle_question.id == vehicle_question.id
                        ),
                        None,
                    )
                    quantity_propio = fleet.quantity_owned if fleet else 0
                    quantity_tercero = fleet.quantity_third_party if fleet else 0
                    quantity_arrendado = fleet.quantity_arrended if fleet else 0
                    quantity_contratista = fleet.quantity_contractors if fleet else 0
                    quantity_intermediacion = (
                        fleet.quantity_intermediation if fleet else 0
                    )
                    quantity_leasing = fleet.quantity_leasing if fleet else 0
                    quantity_renting = fleet.quantity_renting if fleet else 0

                    row_cells[5].text = str(quantity_propio)
                    row_cells[6].text = str(quantity_tercero)
                    row_cells[7].text = str(quantity_arrendado)
                    row_cells[8].text = str(quantity_contratista)
                    row_cells[9].text = str(quantity_intermediacion)
                    row_cells[10].text = str(quantity_leasing)
                    row_cells[11].text = str(quantity_renting)
                    for cell in row_cells:
                        align_cell_text(cell)

                    # Sumar cantidades a los totales
                    total_propio += quantity_propio
                    total_tercero += quantity_tercero
                    total_arrendado += quantity_arrendado
                    total_contratista += quantity_contratista
                    total_intermediacion += quantity_intermediacion
                    total_leasing += quantity_leasing
                    total_renting += quantity_renting

                    total = (
                        total_propio
                        + total_tercero
                        + total_arrendado
                        + total_contratista
                        + total_intermediacion
                        + total_leasing
                        + total_renting
                    )
                    # Agregar fila con los totales
                total_row = table.add_row().cells
                total_row[0].text = "Total Vehiculos".upper()
                total_row[0].merge(total_row[4])
                align_cell_text(total_row[0], "left")
                set_cell_background_color(total_row[0], "2f4858")
                set_cell_text_color(total_row[0])
                total_row[5].text = str(total)
                total_row[5].merge(total_row[6])
                total_row[5].merge(total_row[7])
                total_row[5].merge(total_row[8])
                total_row[5].merge(total_row[9])
                total_row[5].merge(total_row[10])
                total_row[5].merge(total_row[11])

                # Agregar fila para Conductores
                driver_header_row = table.add_row().cells
                driver_header_row[0].text = (
                    "Personas que conducen con fines misionales".upper()
                )
                driver_header_row[0].merge(driver_header_row[8])
                driver_header_row[9].text = "Cantidad"
                driver_header_row[9].merge(driver_header_row[11])
                for cell in driver_header_row:
                    align_cell_text(cell, "left")
                    set_cell_background_color(cell, "2f4858")
                    set_cell_text_color(cell)
                total_conductores = 0
                # Datos de conductores
                for driver_question in driver_questions:
                    driver_data_row = table.add_row().cells
                    driver_data_row[0].text = driver_question.name
                    driver_data_row[0].merge(driver_data_row[8])
                    align_cell_text(driver_data_row[0], "left")
                    set_cell_background_color(driver_data_row[0], "2f4858")
                    set_cell_text_color(driver_data_row[0])
                    driver = next(
                        (
                            f
                            for f in driver_data
                            if f.driver_question.id == driver_question.id
                        ),
                        None,
                    )
                    quantity = driver.quantity if driver else 0
                    driver_data_row[9].text = str(quantity)
                    driver_data_row[9].merge(driver_data_row[11])
                    for cell in driver_data_row:
                        align_cell_text(cell)
                    total_conductores += quantity
                # Agregar fila con los totales
                total_driver_row = table.add_row().cells
                total_driver_row[0].text = "Total Conductores".upper()
                total_driver_row[0].merge(total_driver_row[8])
                align_cell_text(total_driver_row[0], "left")
                set_cell_background_color(total_driver_row[0], "2f4858")
                set_cell_text_color(total_driver_row[0])
                total_driver_row[9].text = str(total_conductores)
                total_driver_row[9].merge(total_driver_row[11])

                # Insertar la tabla en la posición correcta
                paragraph._element.getparent().insert(index + 1, table._element)
                index += 1  # Ajustar el índice para el siguiente elemento

                # Insertar el párrafo vacío para separar la tabla del resumen
                doc.add_paragraph()  # Opcional, para mayor claridad visual
                paragraph._element.getparent().insert(
                    index + 1, doc.paragraphs[-1]._element
                )
                index += 1  # Ajustar el índice para el siguiente elemento

                summary_text = (
                    f"De acuerdo con la información anterior, se identifica que la empresa "
                    f"se encuentra en misionalidad {company.mission.id} | {company.mission.name.upper()} "
                    f"y que cuenta con {total_owned} vehículos propiedad de la empresa y "
                    f"{total_quantity_driver} personas con rol de conductor, por lo tanto, se define "
                    f"que debe diseñar e implementar un plan estratégico de seguridad vial "
                    f"“{count_size.name.upper()}”."
                )
                summary_paragraph = doc.add_paragraph(summary_text)
                summary_paragraph.alignment = 0  # Alinear a la izquierda
                # Insertar el párrafo en el lugar correcto
                paragraph._element.getparent().insert(
                    index + 1, summary_paragraph._element
                )
                index += 1  # Ajustar el índice para el siguiente elemento

                # Insertar el párrafo vacío para separar la tabla del resumen
                doc.add_paragraph()  # Opcional, para mayor claridad visual
                paragraph._element.getparent().insert(
                    index + 1, doc.paragraphs[-1]._element
                )
                index += 1  # Ajustar el índice para el siguiente elemento

            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_results(doc: Document, placeholder: str, filtered_data):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"

            # Estilo de la tabla
            table.autofit = True

            # Estilo de la fila de encabezado
            heading_row = table.rows[0].cells
            heading_row[0].text = "PASO PESV"
            heading_row[1].text = "REQUISITO"
            heading_row[1].merge(heading_row[5])

            for cell in heading_row:
                align_cell_text(cell, "left", "center")

            for data in filtered_data:
                for steps in data["steps"]:
                    step_number = str(steps["step"])
                    for requirement in steps["requirements"]:
                        step_row = table.add_row().cells
                        step_row[0].text = str(steps["step"])  # Paso
                        step_row[1].text = str(requirement["requirement_name"])

                        step_row[1].merge(step_row[5])
                        for cell in step_row:
                            set_cell_background_color(cell, "2f4858")
                            set_cell_text_color(cell)
                        step_row = table.add_row().cells
                        step_row[0].text = "Criterio de verificación"
                        step_row[0].merge(step_row[4])
                        step_row[5].text = "Nivel de Cumplimiento"
                        for cell in step_row:
                            align_cell_text(cell, "left")
                            set_cell_background_color(cell, "ebedf3")
                        question_number = 1
                        for question in requirement["questions"]:
                            question_row = table.add_row().cells
                            question_cell = question_row[0]
                            para = question_cell.add_paragraph()
                            # Run para la numeración en negrita
                            run_number = para.add_run(
                                f"{step_number}.{question_number} "
                            )
                            run_number.bold = True
                            para.add_run(question["question_name"])
                            question_cell.merge(question_row[4])
                            align_cell_text(question_cell, "left")
                            question_row[5].text = question["compliance"]
                            align_cell_text(question_row[5])
                            question_number += 1
            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_recomendations(
    doc: Document, placeholder: str, recomendaciones_agrupadas
):
    VALID_STEPS = {
        "P": "PLANEAR",
        "H": "HACER",
        "V": "VERIFICAR",
        "A": "ACTUAR",
    }
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"

            # Estilo de la tabla
            table.autofit = True
            # Encabezado de la tabla
            heading_row = table.rows[0].cells
            heading_row[0].text = "CICLO PESV"
            heading_row[0].merge(heading_row[5])

            # Insertar datos por ciclo
            for item in recomendaciones_agrupadas:

                ciclo = VALID_STEPS.get(item["cycle"].upper(), "Otros").upper()
                recomendaciones = item["recomendations"]
                row = table.add_row().cells
                row[0].text = ciclo
                row[0].merge(row[5])
                align_cell_text(row[0])
                if item["cycle"].upper() == "P":
                    set_cell_background_color(row[0], "0066B2")
                elif item["cycle"].upper() == "H":
                    set_cell_background_color(row[0], "00A551")
                elif item["cycle"].upper() == "V":
                    set_cell_background_color(row[0], "DCB00A")
                elif item["cycle"].upper() == "A":
                    set_cell_background_color(row[0], "EC1C24")
                set_cell_text_color(row[0])

                row = table.add_row().cells
                cell = row[0]
                cell.paragraphs[0].clear()  # Limpiar cualquier texto existente

                for recomendacion in recomendaciones:
                    p = cell.add_paragraph(recomendacion["recomendacion"])
                    apply_bullets(p)
                    if recomendacion["observation"]:
                        p2 = cell.add_paragraph(recomendacion["observation"])
                        apply_bullets(p2)  # Aplicar viñetas al segundo párrafo
                row[0].merge(row[5])
                for cell in row:
                    bottom_border = OxmlElement("w:bottom")
                    bottom_border.set(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                        "single",
                    )
                    cell._element.get_or_add_tcPr().append(bottom_border)

            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def set_vertical_cell_direction(cell: _Cell, direction: str, align_center: bool = True):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement("w:textDirection")
    textDirection.set(qn("w:val"), direction)  # btLr tbRl
    tcPr.append(textDirection)
    # Configuración de la alineación vertical
    if align_center:
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)


def merge_cells_vertically(cell):
    """Combine vertical cells in the given cell."""
    cell._tc.get_or_add_tcPr().append(OxmlElement("w:vMerge"))


def insert_table_work_plan(doc: Document, placeholder: str, data: dict):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "PLAN DE TRABAJO"
            hdr_cells[0].merge(hdr_cells[5])
            hdr_cells[6].text = "HORAS"
            hdr_cells[7].text = ""  # Empty cell
            # Agregar datos a la tabla
            VALID_STEPS = {
                "P": "PLANEAR",
                "H": "HACER",
                "V": "VERIFICAR",
                "A": "ACTUAR",
            }
            for cycle, recommendations in data.items():
                # Fila para el ciclo
                phase_name = VALID_STEPS.get(cycle.upper(), "Otros").upper()
                cycle_row = table.add_row().cells
                cycle_row[0].text = phase_name
                cycle_row[0].merge(cycle_row[7])
                align_cell_text(cycle_row[0])
                if cycle.upper() == "P":
                    set_cell_background_color(cycle_row[0], "0066B2")
                elif cycle.upper() == "H":
                    set_cell_background_color(cycle_row[0], "00A551")
                elif cycle.upper() == "V":
                    set_cell_background_color(cycle_row[0], "DCB00A")
                elif cycle.upper() == "A":
                    set_cell_background_color(cycle_row[0], "EC1C24")
                set_cell_text_color(cycle_row[0])
                # Añadir filas para las recomendaciones bajo el ciclo
                for rec in recommendations:
                    rec_row = table.add_row().cells
                    rec_row[0].text = rec["recommendation_name"]
                    rec_row[0].merge(rec_row[5])
            total_row = table.add_row().cells
            total_row[0].text = "TOTAL HORAS"
            total_row[0].merge(total_row[5])
            set_cell_background_color(total_row[0], "2f4858")
            set_cell_background_color(total_row[6], "2f4858")
            set_cell_background_color(total_row[7], "2f4858")
            set_cell_text_color(total_row[0])
            set_cell_text_color(total_row[6])
            set_cell_text_color(total_row[7])
            align_cell_text(total_row[0])

            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_conclusion(
    doc: Document, placeholder: str, datas_by_cycle, sizeName: str
):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"

            heading_row = table.rows[0].cells
            heading_row[0].text = (
                f"ESTRUCTURA DE PONDERACIÓN \n NIVEL {sizeName} - PESV"
            )
            heading_row[0].merge(heading_row[7])
            align_cell_text(heading_row[0], "center", "center")

            title_row = table.add_row().cells
            title_row[0].text = "FASE"
            title_row[1].text = "PASO"
            title_row[2].text = "DESCRIPCIÓN"
            title_row[2].merge(title_row[5])
            title_row[6].text = "% PASO"
            title_row[7].text = "% FASE"
            for cell in title_row:
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)

            VALID_STEPS = {
                "P": "PLANEAR",
                "H": "HACER",
                "V": "VERIFICAR",
                "A": "ACTUAR",
            }

            # # Recorrer y agregar filas para cada fase
            for cycle in datas_by_cycle:
                phase_name = VALID_STEPS.get(cycle["cycle"].upper(), "Otros").upper()
                for i, requerimiento in enumerate(cycle["steps"]):
                    cells = table.add_row().cells
                    if i == 0:
                        percentage = round(cycle["cycle_percentage"], 2)
                        set_vertical_cell_direction(cells[0], "tbRl")
                        cells[0].text = phase_name
                        align_cell_text(cells[0], "center", "center")
                        if cycle["cycle"].upper() == "P":
                            set_cell_background_color(cells[0], "0066B2")
                        elif cycle["cycle"].upper() == "H":
                            set_cell_background_color(cells[0], "00A551")
                        elif cycle["cycle"].upper() == "V":
                            set_cell_background_color(cells[0], "DCB00A")
                        elif cycle["cycle"].upper() == "A":
                            set_cell_background_color(cells[0], "EC1C24")
                        set_cell_text_color(cells[0])

                        cells[7].text = f"{percentage}%"
                        align_cell_text(cells[7], "center", "center")

                    cells[1].text = str(requerimiento["step"])
                    for requirement in requerimiento["requirements"]:
                        cells[2].text = str(requirement["requirement_name"])
                        cells[2].merge(cells[5])
                    req_percentage = round(requerimiento["percentage"], 2)
                    cells[6].text = f"{req_percentage}%"
                    align_cell_text(cells[6], "center", "center")

                start_idx = len(table.rows) - len(cycle["steps"])
                end_idx = len(table.rows) - 1
                for row_idx in range(start_idx, end_idx + 1):
                    table.cell(row_idx, 0).merge(table.cell(end_idx, 0))
                    table.cell(row_idx, 7).merge(table.cell(end_idx, 7))

            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_conclusion_articulated(
    doc: Document, placeholder: str, datas_by_cycle, sizeName: str
):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"
            heading_row = table.rows[0].cells
            heading_row[0].text = (
                f"ESTRUCTURA DE PONDERACIÓN \n NIVEL {sizeName} - PESV - ARTICULACIONES"
            )
            heading_row[0].merge(heading_row[7])

            align_cell_text(heading_row[0], "center", "center")
            title_row = table.add_row().cells
            title_row[0].text = "PASO"
            title_row[1].text = "DESCRIPCIÓN"
            title_row[1].merge(title_row[6])
            title_row[7].text = "NIVEL CUMPLIMIENTO"
            align_cell_text(title_row[7], "center", "center")
            for cell in title_row:
                align_cell_text(cell, "left", "center")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)

            VALID_STEPS = {
                "P": "PLANEAR",
                "H": "HACER",
                "V": "VERIFICAR",
                "A": "ACTUAR",
            }
            # To track rows for merging compliance column
            row_indices = []
            for cycle_data in datas_by_cycle:
                phase_name = VALID_STEPS.get(
                    cycle_data["cycle"].upper(), "Otros"
                ).upper()

                # Add a row for the phase name
                phase_row = table.add_row().cells
                phase_row[0].text = phase_name
                phase_row[0].merge(phase_row[7])
                align_cell_text(phase_row[0], "center", "center")
                if cycle_data["cycle"].upper() == "P":
                    set_cell_background_color(phase_row[0], "0066B2")
                elif cycle_data["cycle"].upper() == "H":
                    set_cell_background_color(phase_row[0], "00A551")
                elif cycle_data["cycle"].upper() == "V":
                    set_cell_background_color(phase_row[0], "DCB00A")
                elif cycle_data["cycle"].upper() == "A":
                    set_cell_background_color(phase_row[0], "EC1C24")
                set_cell_text_color(phase_row[0])

                for requerimiento in cycle_data["steps"]:
                    step_start_row_index = len(table.rows)
                    # Add a row for each step-requirement
                    for requirement_data in requerimiento["requirements"]:
                        requirement_name = requirement_data["requirement_name"]
                        last_question = requirement_data["questions"][-1]
                        last_question_name = last_question["question_name"]
                        last_compliance = last_question["compliance"]

                        # Add row for step and requirement
                        step_row = table.add_row().cells
                        step_row[0].text = (
                            f"{requerimiento['step']} - {requirement_name}"
                        )
                        step_row[0].merge(step_row[6])  # Merge description columns
                        step_row[7].text = last_compliance
                        align_cell_text(step_row[7], "center", "center")

                        # Add row for last question
                        question_row = table.add_row().cells
                        question_row[0].text = ""  # Empty cell for FASE
                        question_row[1].text = last_question_name
                        question_row[1].merge(
                            question_row[6]
                        )  # Merge description columns
                        question_row[7].text = ""  # Empty compliance cell

                    row_indices.append((step_start_row_index, len(table.rows) - 1))
            # Merge "NIVEL CUMPLIMIENTO" cells vertically
            for start_idx, end_idx in row_indices:
                if end_idx < len(table.rows) and start_idx < len(table.rows):
                    cell = table.cell(start_idx, 7)
                    cell.text = cell.text  # Ensure cell has text
                    for idx in range(start_idx + 1, end_idx + 1):
                        table.cell(idx, 7).text = ""  # Clear text in cells to be merged
                    cell.merge(table.cell(end_idx, 7))
                    trim_merged_cell(cell)  # Trim merged cell
                    align_cell_text(cell, "center", "center")  # Align merged cell
            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_conclusion_percentage_articuled(
    doc: Document, placeholder: str, datas_by_cycle
):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            heading_row = table.rows[0].cells
            heading_row[0].text = "PASOS EVALUADOS"
            heading_row[0].merge(heading_row[1])
            heading_row[2].text = "SIN ARTICULAR"
            heading_row[2].merge(heading_row[3])
            for cell in heading_row:
                align_cell_text(cell, "left", "center")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)
            count_cumple = 0
            count_no_cumple = 0
            count_cumple_parcial = 0
            count_no_aplica = 0
            for cycle_data in datas_by_cycle:
                for step_data in cycle_data["steps"]:
                    for requirement_data in step_data["requirements"]:
                        # Obtener la última pregunta
                        last_question = requirement_data["questions"][-1]
                        last_compliance = last_question["compliance"]

                        if last_compliance == "CUMPLE":
                            count_cumple += 1
                        elif last_compliance == "NO CUMPLE":
                            count_no_cumple += 1
                        elif last_compliance == "CUMPLE PARCIALMENTE":
                            count_cumple_parcial += 1
                        elif last_compliance == "NO APLICA":
                            count_no_aplica += 1

            title_row = table.add_row().cells
            title_row[0].text = str(
                count_cumple + count_no_cumple + count_cumple_parcial + count_no_aplica
            )
            title_row[0].merge(title_row[1])
            align_cell_text(title_row[0], "center", "center")

            # title_row[1].text = str(count_cumple)  # Cumple
            title_row[2].text = str(count_no_cumple)  # No Cumple
            title_row[2].merge(title_row[3])
            align_cell_text(title_row[2], "center", "center")
            # title_row[3].text = str(count_cumple_parcial)  # Cumple Parcialmente
            # title_row[4].text = str(count_no_aplica)  # No Aplica

            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_table_conclusion_percentage(
    doc: Document, placeholder: str, counts, perecentaje
):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            index = paragraph._element.getparent().index(paragraph._element)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            heading_row = table.rows[0].cells
            heading_row[0].text = "TOTAL ITEMS"
            heading_row[1].text = "CUMPLE"
            heading_row[2].text = "NO CUMPLE"
            heading_row[3].text = "CUMPLE PARCIALMENTE"
            heading_row[4].text = "NO APLICA"
            heading_row[5].text = "PORCENTJAE CUMPLIMIENTO"
            for cell in heading_row:
                align_cell_text(cell, "left", "center")
                set_cell_background_color(cell, "2f4858")
                set_cell_text_color(cell)

            title_row = table.add_row().cells

            for compliance in counts:
                compliance_id = compliance.id
                count = compliance.count if compliance.count is not None else 0

                # Asigna el valor en la celda correspondiente en tu tabla
                if compliance_id == 1:  # Cumple
                    title_row[1].text = str(count)
                elif compliance_id == 2:  # No Cumple
                    title_row[2].text = str(count)
                elif compliance_id == 3:  # Cumple Parcialmente
                    title_row[3].text = str(count)
                elif compliance_id == 4:  # No Aplica
                    title_row[4].text = str(count)
            total_items = sum(int(title_row[i].text) for i in range(1, 5))
            title_row[0].text = str(total_items)
            title_row[5].text = f"{perecentaje}%"
            table._element.getparent().insert(index + 1, table._element)
            return  # Salir después de insertar la tabla para evitar múltiples inserciones


def insert_image_after_placeholder(doc, placeholder, image_path):
    # Iterar sobre todos los párrafos en el documento
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Crear un nuevo párrafo para la imagen
            new_paragraph = para.insert_paragraph_before()
            run = new_paragraph.add_run()
            run.add_picture(image_path, width=Inches(5))
            # Eliminar el párrafo original con el placeholder
            para.clear()  # Eliminar el texto del párrafo pero mantener el párrafo
            break


def create_radar_chart(datas_by_cycle):
    labels = ["PLANEAR", "HACER", "ACTUAR", "VERIFICAR"]  # 5 categorías
    stats = []  # Valores para cada categoría
    for item in datas_by_cycle:
        stats.append([round(item["cycle_percentage"], 2)])
    # Número de categorías
    num_vars = len(labels)

    # Ángulos para cada categoría
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()

    # Añadir el primer ángulo al final para cerrar el gráfico
    stats += stats[:1]
    angles += angles[:1]

    # Crear la gráfica de telaraña
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, stats, color="blue", alpha=0.25)  # Rellenar el área bajo la curva

    # Configurar los límites radiales
    ax.set_ylim(0, 100)

    # Añadir etiquetas
    ax.set_thetagrids(
        np.degrees(angles[:-1]), labels
    )  # Usar angles[:-1] para evitar duplicar la última etiqueta

    # Guardar como imagen PNG
    image_file = "telarana.png"
    plt.savefig(image_file, bbox_inches="tight")

    # Mostrar la gráfica (opcional, según se necesite)
    # plt.show()

    return image_file


def create_bar_chart(datas_by_cycle):
    wb = openpyxl.Workbook()
    ws = wb.active
    data = [["Paso PESV", "Porcentage"]]
    for item in datas_by_cycle:
        for steps in item["steps"]:
            step = steps["step"]
            percentage = steps["percentage"]
            data.append([step, percentage])
    for row in data:
        ws.append(row)
    chart = BarChart()
    chart.title = "NIVEL DEL CUMPLIMIENTO DEL PESV"
    chart.x_axis.title = "Paso PESV"
    chart.y_axis.title = "Porcentage"
    data = Reference(ws, min_col=2, min_row=1, max_col=2, max_row=len(data))
    categories = Reference(ws, min_col=1, min_row=2, max_row=len(data))
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)
    ws.add_chart(chart, "E5")
    excel_file = "chart.xlsx"
    wb.save(excel_file)

    df = pd.read_excel(excel_file)
    fig, ax = plt.subplots()
    df.plot(kind="bar", x="Paso PESV", y="Porcentage", ax=ax)
    ax.set_title("NIVEL DEL CUMPLIMIENTO DEL PESV")
    ax.set_xlabel("Paso PESV")
    ax.set_ylabel("Porcentage")
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    img_buffer.seek(0)
    image_file = "chart.png"
    with open(image_file, "wb") as f:
        f.write(img_buffer.getvalue())

    return image_file


def convert_docx_to_pdf_base64(docx_bytes: bytes) -> str:

    with NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx_file:
        temp_docx_file.write(docx_bytes)
        temp_docx_name = temp_docx_file.name

    with NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf_file:
        temp_pdf_name = temp_pdf_file.name

    try:
        convert(temp_docx_name, temp_pdf_name)

        with open(temp_pdf_name, "rb") as pdf_file:
            pdf_content = pdf_file.read()
            pdf_base64 = base64.b64encode(pdf_content).decode("utf-8")
    finally:
        # Ensure the files are cleaned up even if an error occurs
        os.remove(temp_docx_name)
        os.remove(temp_pdf_name)

    return pdf_base64


def calculate_obtained_value(num_questions):
    if num_questions == 0:
        return 0
    # Supongamos que cada pregunta vale un punto
    return num_questions
